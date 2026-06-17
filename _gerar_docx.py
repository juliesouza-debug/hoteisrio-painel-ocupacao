# -*- coding: utf-8 -*-
"""Gera as versões Word (.docx) da Nota Técnica e do Relatório."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
BLUE = RGBColor(0x25,0x63,0xeb)
NAVY = RGBColor(0x0f,0x1d,0x33)
ORANGE = RGBColor(0xc2,0x41,0x0c)
GREY = RGBColor(0x5b,0x6b,0x85)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),hexc)
    tcPr.append(shd)

def runs_from(p, text, base_bold=False, color=None, size=None):
    """Interpreta **negrito** dentro do texto."""
    parts = text.split("**")
    for i,seg in enumerate(parts):
        if seg=="": continue
        r = p.add_run(seg)
        r.bold = base_bold or (i % 2 == 1)
        if color is not None: r.font.color.rgb = color
        if size is not None: r.font.size = Pt(size)

def para(doc, text, bold=False, color=None, size=10.5, space=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if align: p.alignment = align
    runs_from(p, text, base_bold=bold, color=color, size=size)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    runs_from(p, text, size=10.5)
    return p

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'2563eb')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def box(doc, lines, fill, title=None):
    t = doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0,0); shade(cell, fill)
    cell.paragraphs[0].text=""
    if title:
        p=cell.paragraphs[0]; r=p.add_run(title); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY
    for ln in lines:
        p = cell.add_paragraph(); p.paragraph_format.space_after=Pt(3)
        runs_from(p, ln, size=10.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,htext in enumerate(header):
        c=t.rows[0].cells[j]; shade(c,'F4F6FB'); c.paragraphs[0].text=""
        r=c.paragraphs[0].add_run(htext); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    for row in rows:
        cells=t.add_row().cells
        for j,val in enumerate(row):
            cells[j].paragraphs[0].text=""
            runs_from(cells[j].paragraphs[0], str(val), size=10)
    return t

def header_block(doc, kicker, title, subtitle):
    # faixa com logo
    logo = os.path.join(BASE,"logo_hoteisrio_azul.png")
    if os.path.exists(logo):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(logo, height=Inches(0.62))
        p.paragraph_format.space_after=Pt(2)
    if kicker:
        p=para(doc, kicker.upper(), bold=True, color=ORANGE, size=9, space=1)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(title); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=NAVY
    para(doc, subtitle, color=GREY, size=10, space=10)
    # linha
    sep=doc.add_paragraph(); pPr=sep._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'18'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'f97316')
    pbdr.append(bottom); pPr.append(pbdr); sep.paragraph_format.space_after=Pt(10)

def base_doc():
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
    for s in doc.sections:
        s.top_margin=Inches(0.7); s.bottom_margin=Inches(0.7); s.left_margin=Inches(0.85); s.right_margin=Inches(0.85)
    return doc

# =================== NOTA TÉCNICA ===================
def gerar_nota():
    doc=base_doc()
    header_block(doc, "Nota Técnica · Posicionamento Institucional",
        "A necessidade de regulamentação adequada do aluguel de temporada (STR)",
        "HotéisRIO — Sindicato de Hotéis e Meios de Hospedagem do Município do Rio de Janeiro · "
        "Subsídio para diálogo com o poder público · Junho/2026")

    box(doc, [
        "O aluguel de temporada exercido de forma **profissional, reiterada e em escala** não é locação: é "
        "**prestação de serviço de hospedagem**. Tratá-lo sob a Lei do Inquilinato é um **enquadramento jurídico "
        "equivocado** que cria um mercado de hotelaria informal — sem tributação, sem fiscalização sanitária, sem "
        "segurança contra incêndio e sem responsabilidade trabalhista equivalentes. Não se pede proibição: pede-se "
        "**isonomia e regulamentação**, em defesa do consumidor, do erário e da cidade."
    ], 'E8EEFB', title="Tese central")

    h1(doc,"1. O argumento jurídico: hospedagem não é locação")
    para(doc,"A distinção é de natureza, e não de duração:")
    table(doc, ["Locação por temporada (Lei 8.245/91, art. 48)","STR profissional (o que ocorre de fato)"],
        [["Cessão do uso de um imóvel a um locatário determinado","Oferta ao público em plataforma, contínua e impessoal"],
         ["Relação bilateral e pessoal entre locador e inquilino","Hóspedes rotativos, diárias, check-in/out, limpeza"],
         ["Sem prestação de serviços; sem oferta pública contínua","Operadores com múltiplas unidades e gestão profissional"],
         ["Prazo até 90 dias para fim específico (lazer, curso, saúde)","Em essência, um meio de hospedagem (Lei 11.771/2008)"]])
    box(doc, [
        "**Precedente que sustenta a tese:** o Superior Tribunal de Justiça já reconheceu (REsp 1.819.075/RS, 4ª Turma, "
        "2021) que a locação de curtíssima temporada via plataformas é uma **“nova modalidade de hospedagem atípica”** — "
        "e não locação residencial típica —, admitindo inclusive que condomínios residenciais a restrinjam.",
        "(Referências legais e jurisprudenciais a serem confirmadas/atualizadas pela assessoria jurídica antes do uso oficial.)"
    ], 'EFF5FF')
    para(doc,"Conclusão: o que define a atividade é a **natureza hoteleira** (serviço, rotatividade, oferta pública, "
        "profissionalismo), e não o rótulo de “temporada”. O enquadramento na Lei do Inquilinato serve, na prática, "
        "para **fugir das obrigações da hotelaria**.")

    h1(doc,"2. A escala: uma rede hoteleira paralela")
    para(doc,"O fenômeno deixou de ser doméstico. Os dados da cidade mostram operação de porte hoteleiro:")
    bullet(doc,"**~29.482 imóveis** de temporada com reservas por mês (2025);")
    bullet(doc,"**+74%** de crescimento da oferta em apenas dois anos;")
    bullet(doc,"**3,1×** mais imóveis de temporada do que quartos de hotel em Ipanema/Leblon.")
    para(doc,"Oferta equivalente às **31.644 UHs** da hotelaria formal, operando majoritariamente fora do marco "
        "regulatório. O alvo não é o cidadão que aluga esporadicamente, e sim o **operador profissional** que mantém "
        "dezenas de unidades em regime de hotel.")

    h1(doc,"3. Isonomia concorrencial e tributária")
    para(doc,"Quem hospeda deve cumprir as mesmas regras de quem hospeda. Hoje há assimetria que distorce a "
        "concorrência e subtrai receita pública:")
    table(doc, ["Obrigação","Hotel","STR profissional (em regra)"],
        [["Cadastro turístico (Cadastur)","Sim","Não exigido / não fiscalizado"],
         ["ISS sobre hospedagem + tributos municipais","Sim","Em geral não recolhe ISS"],
         ["Licença sanitária / vigilância","Sim","Não"],
         ["Auto de Vistoria do Corpo de Bombeiros (AVCB)","Sim","Não"],
         ["Alvará de funcionamento / acessibilidade","Sim","Não"],
         ["Vínculos trabalhistas (CLT) e encargos","Sim","Informal / terceirizado"]])
    para(doc,"Resultado: **concorrência desigual** (custo regulatório só de um lado) e **perda de arrecadação** "
        "municipal de ISS e taxas — recursos que financiam a própria cidade que recebe o turista.")

    h1(doc,"4. Segurança e defesa do consumidor")
    bullet(doc,"**Segurança contra incêndio:** imóveis sem AVCB, sem brigada, sinalização ou extintores — risco à vida.")
    bullet(doc,"**Vigilância sanitária:** ausência de controle de higiene, potabilidade e pragas.")
    bullet(doc,"**Segurança pública:** hóspedes não identificados em prédios residenciais, sem registro.")
    bullet(doc,"**Defesa do consumidor:** proteção e responsabilização inferiores às do hóspede de hotel.")

    h1(doc,"5. Impacto habitacional e urbano")
    para(doc,"A conversão de moradias em hospedagem permanente **retira imóveis do estoque residencial**, pressiona o "
        "aluguel de longo prazo e acelera a “turistificação”. Foi o que levou **Barcelona, Lisboa, Paris, Amsterdã e "
        "Nova York** a licenciar, cadastrar e limitar dias de operação do STR. Regular é a **regra internacional** — e "
        "protege o morador, não só o hoteleiro.")

    h1(doc,"6. O que se propõe (não é proibir, é regular)")
    box(doc, [
        "• **Classificação correta:** reconhecer o STR profissional como **meio de hospedagem** (Lei Geral do Turismo), e não como locação.",
        "• **Cadastro obrigatório** de imóveis e operadores, com número visível no anúncio.",
        "• **Isonomia tributária:** ISS e taxas equivalentes às da hotelaria.",
        "• **Exigências mínimas de segurança** (incêndio/sanitária) proporcionais à atividade.",
        "• **Responsabilidade das plataformas** por anunciar apenas imóveis regularizados e repassar dados ao fisco.",
        "• **Limites por zona/condomínio**, preservando o uso residencial onde couber.",
    ], 'EFFDF4')

    h1(doc,"7. Antecipando os contra-argumentos")
    table(doc, ["Dirão que…","Resposta"],
        [["“Regular é proibir / atrasar a inovação.”","Não se pede proibição, e sim cadastro e isonomia — como já fazem as principais cidades turísticas do mundo."],
         ["“O STR só democratiza renda do pequeno proprietário.”","O ocasional pode ter regra simplificada. O alvo é o operador profissional com múltiplas unidades — hotel sem se chamar hotel."],
         ["“O hotel só quer eliminar concorrência.”","Os dados mostram que a ocupação hoteleira cresceu mesmo com a expansão do STR. O pleito é por regras iguais, segurança e arrecadação."],
         ["“Falta base legal.”","O STJ já caracterizou a atividade como hospedagem atípica; a Lei Geral do Turismo já define meios de hospedagem. A lacuna é de fiscalização e regulamentação municipal."]])

    h1(doc,"8. Mensagem de encerramento")
    para(doc,"A defesa não é dos hotéis contra a tecnologia — é do **interesse público**: que toda hospedagem, qualquer "
        "que seja a plataforma, pague seus tributos, garanta a segurança do hóspede e respeite a cidade e seus "
        "moradores. **Mesma atividade, mesmas regras.**", size=11)

    para(doc,"Documento de subsídio elaborado a partir do painel e do relatório de ocupação HotéisRIO × AirDNA "
        "(triênio 2023–2025). As referências legais (Lei 8.245/1991; Lei 11.771/2008; REsp 1.819.075/RS) devem ser "
        "validadas pela assessoria jurídica antes de apresentação oficial.", color=GREY, size=8.5, space=2)

    out=os.path.join(BASE,"Nota_Tecnica_Regulamentacao_Temporada.docx"); doc.save(out); return out

# =================== RELATÓRIO ===================
def gerar_relatorio():
    doc=base_doc()
    header_block(doc, None,
        "Impacto do Aluguel de Temporada sobre a Hotelaria",
        "Cidade do Rio de Janeiro · Análise do triênio 2023–2025 · Fontes: pesquisa HotéisRIO & AirDNA · "
        "Documento de apoio técnico · Junho/2026")

    box(doc, [
        "Entre 2023 e 2025, o aluguel de temporada consolidou-se como uma **oferta paralela de hospedagem de porte "
        "equivalente ao da hotelaria**. Ainda assim, os dados mostram mais **coexistência em um mercado em expansão** "
        "do que substituição direta:",
        "• A oferta de temporada **cresceu 74%** em dois anos (~16,9 mil → ~29,5 mil/mês), aproximando-se das **31.644 UHs** da hotelaria.",
        "• Mesmo assim, a **ocupação dos hotéis subiu** (71,1% → 77,7%) e a **vantagem sobre a temporada aumentou** (de 9,9 para 17,6 pontos). Não há canibalização agregada.",
        "• A pressão é **concentrada e sazonal**: forte em bairros residenciais e nos picos de eventos, quando a temporada pratica diárias de até R$ 613.",
    ], 'F4F6FB', title="Resumo executivo")

    h1(doc,"1. O tamanho do fenômeno")
    para(doc,"A oferta de temporada já rivaliza em volume com a hotelaria (unidades não diretamente comparáveis — ver observação):")
    bullet(doc,"**31.644** quartos (UHs) em 286 hotéis;")
    bullet(doc,"**~29.482** imóveis de temporada com reservas/mês (2025);")
    bullet(doc,"**+74%** de crescimento da oferta de temporada 2023→2025 (mais de 12 mil imóveis em dois anos).")

    h1(doc,"2. Impactos negativos para a hotelaria")
    para(doc,"**2.1 Oferta concorrente de grande escala e crescimento acelerado.** Estoque de hospedagem de porte "
        "semelhante ao hoteleiro, em expansão muito mais rápida e com menor exigência regulatória/tributária — o que "
        "limita o poder de elevar preços e ocupação.")
    para(doc,"**2.2 Penetração elevada em bairros-chave.** Onde o hóspede de lazer se concentra, a temporada já supera "
        "a hotelaria em unidades:")
    table(doc, ["Região","Hotéis","UHs","Imóveis temporada*","Temp ÷ UHs"],
        [["Ipanema / Leblon","18","1.372","4.278","3,1×"],
         ["Leme / Copacabana","87","10.970","11.066","1,0×"],
         ["Centro","71","5.940","4.572","0,8×"],
         ["São Conrado / Barra / Recreio","44","8.390","3.666","0,4×"],
         ["Flamengo / Botafogo","66","4.972","2.046","0,4×"],
         ["**Cidade**","**286**","**31.644**","**29.482**","**0,9×**"]])
    para(doc,"*Média mensal de imóveis com reservas em 2025.", color=GREY, size=9, space=6)
    para(doc,"**2.3 Disputa agressiva nos picos.** Nos eventos de melhor margem (Réveillon, Carnaval, verão), a "
        "temporada pratica suas diárias mais altas: de ~R$ 500 (dez) para **R$ 613 em fev/2026**, contra ~R$ 375 na "
        "baixa estação — reduzindo a escassez que sustentaria tarifas hoteleiras maiores nessas datas.")

    h1(doc,"3. Impactos positivos e fatores de resiliência")
    para(doc,"**3.1 Sem canibalização agregada.** Apesar do salto de 74% na oferta de temporada, a ocupação dos hotéis "
        "subiu e a vantagem quase dobrou:")
    table(doc, ["Ocupação média","2023","2024","2025"],
        [["Hotéis","71,1%","73,2%","77,7%"],
         ["Temporada","61,2%","56,9%","60,1%"],
         ["**Vantagem da hotelaria**","**+9,9 pp**","**+16,3 pp**","**+17,6 pp**"]])
    para(doc,"Leitura: **expansão da demanda total** pelo destino, absorvida pelas duas modalidades — não transferência de hóspedes.")
    para(doc,"**3.2 Hotelaria ocupa mais e com mais estabilidade** (amplitude sazonal de 19 pontos em 2025, contra 26 da "
        "temporada): demanda corporativa e de eventos, padronização, segurança e previsibilidade.")
    para(doc,"**3.3 Absorção de excedente em megaeventos** protege o destino, evitando que o Rio “venda esgotado”.")
    para(doc,"**3.4 Validação de diárias mais altas** (ADR de R$ 361 para R$ 438; RevPAR de R$ 221 para R$ 264), o que "
        "ancora para cima a percepção de preço da cidade.")

    h1(doc,"4. Leitura por região (síntese)")
    bullet(doc,"**Ipanema / Leblon** — maior pressão: temporada 3× os quartos de hotel; competir por diferenciação, não preço.")
    bullet(doc,"**Leme / Copacabana** — equilíbrio (1:1) e maior volume absoluto; mercado mais disputado em receita.")
    bullet(doc,"**Centro** — hotelaria predominante (perfil corporativo); temporada mais frágil na baixa.")
    bullet(doc,"**Barra/Recreio/São Conrado** e **Flamengo/Botafogo** — hotelaria dominante (~0,4×); menor exposição.")

    h1(doc,"5. Conclusão e recomendações")
    para(doc,"O aluguel de temporada é, ao mesmo tempo, concorrente relevante e componente da capacidade turística. No "
        "triênio, não comprimiu a ocupação hoteleira, mas pressiona margens em segmentos e datas e avança com menor "
        "regulação. Recomenda-se:")
    bullet(doc,"**Atuação institucional** por isonomia regulatória e tributária, sustentada por estes indicadores.")
    bullet(doc,"**Estratégia comercial por região e sazonalidade**: defender diária nos picos; reforçar corporativo/eventos na baixa.")
    bullet(doc,"**Monitoramento contínuo** da relação imóveis ÷ UHs e do diferencial de ocupação por bairro, via painel interativo.")

    h1(doc,"6. Metodologia e limitações")
    bullet(doc,"**Fontes:** ocupação HotéisRIO (5 regiões, mensal); desempenho da temporada do AirDNA (20 bairros, ponderados por nº de imóveis).")
    bullet(doc,"**Período:** dados de temporada a partir de maio/2023.")
    bullet(doc,"**Inventário não comparável:** hotel conta **quartos (UHs duplos)**; temporada conta **imóveis** (Airbnb pode ter +6 quartos). O nº de imóveis refere-se aos com reservas no mês.")
    bullet(doc,"**A hotelaria não coleta diária/RevPAR** nesta base — comparação de preço/receita restrita à temporada.")
    bullet(doc,"**Causalidade:** correlação não prova causa; fatores externos (retomada pós-pandemia, câmbio, eventos) influenciam as duas modalidades.")

    out=os.path.join(BASE,"Relatorio_Impacto_Temporada_x_Hotelaria.docx"); doc.save(out); return out

print("Nota:", gerar_nota())
print("Relatório:", gerar_relatorio())
